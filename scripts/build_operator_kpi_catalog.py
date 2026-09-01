#!/usr/bin/env python3
"""Build shareable Operator KPI catalogue (Word + PDF)."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.colors import HexColor
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas

OUT_DOCX = "/workspace/operator-kpi-catalog.docx"
OUT_PDF = "/workspace/operator-kpi-catalog.pdf"

SECTIONS = [
    ("Executive north-stars (12–15 KPIs)", [
        "Net sales / revenue growth (by segment, channel, product)",
        "Service margin / EBITDA contribution",
        "Active subscriber base & net adds",
        "Voluntary & involuntary churn; MNP port-in / port-out (operator-wise)",
        "ARPU / ARPA",
        "NPS / CES / complaint rate",
        "Customer-facing network & application availability",
        "P1/P2 MTTR",
        "Order-to-activate lead time",
        "Bill accuracy & billing dispute rate",
        "DSO & collection rate",
        "CAPEX/OPEX per subscriber",
        "Automation / auto-remediation rate (AO maturity)",
    ]),
    ("Commercial & sales (COP · CHQ · digital · partners)", [
        "Gross sales by channel: COP (company-owned stores), CHQ (contact centre), digital, partners",
        "Conversion: lead → quote → order → activation",
        "Attach / upsell rate; promo uptake vs margin",
        "Order fallout & re-quote rate",
        "CAC by channel; partner commission accuracy",
    ]),
    ("Customer lifecycle & retention", [
        "Gross churn, net churn, voluntary vs involuntary",
        "MNP port-out / port-in by competitor",
        "Churn by segment, tenure cohort, high-value base",
        "Early-life churn (0–90 days); save rate; win-back rate",
        "Contract renewal rate",
    ]),
    ("BSS — order-to-cash & revenue cycle", [
        "Order capture volume & lead time (order created → accepted)",
        "Order-to-activate (O2A) lead time; provisioning first-time-right",
        "Port-in / port-out completion time",
        "Billing end-to-end completion time (cycle close T+n)",
        "Bill prep → confirmation time",
        "Billing reruns (count, root cause, cost)",
        "Bill accuracy; rating/mediation errors; revenue leakage",
        "Accounts ready for collection",
        "Accounts in collection (by stage / step clarity)",
        "Accounts pending write-off trigger",
        "DSO; bad debt / write-off rate; dunning effectiveness",
    ]),
    ("OSS / IT application & service management", [
        "Application availability by tier (gold / silver / bronze)",
        "P1/P2 incident volume, duration, MTTR",
        "Repeat incident rate; major incidents",
        "Problem backlog; known-error sunset rate",
        "Change success rate; emergency change rate; failed deployments",
        "SLA attainment by tower (BSS, OSS, CRM, ERP, integration)",
        "API / transaction success rate on critical journeys",
        "Batch job SLA (billing, mediation, reporting)",
        "CMDB accuracy; license utilization",
    ]),
    ("Network (RAN · core · transport · fixed)", [
        "Network availability (2G/3G/4G/5G, fixed broadband)",
        "Call setup success; drop call rate; data session success",
        "Throughput, latency, congestion",
        "VoLTE / VoWiFi success",
        "Site/cell outage count & duration; MTTR (network)",
        "RAN/core utilization & capacity headroom",
    ]),
    ("Infrastructure, cloud & platform", [
        "Compute/storage availability; patch compliance",
        "Cloud spend vs budget (FinOps); unit cost per subscriber/transaction",
        "Backup/restore success; capacity forecast accuracy",
        "Database & middleware performance",
        "Observability coverage (% services with SLO monitoring)",
        "ETL / data pipeline SLA",
    ]),
    ("Security, risk & compliance", [
        "Security incidents; MTTD / MTTC / MTTR",
        "Critical vulnerability backlog & patch SLA",
        "Fraud / access anomalies",
        "Regulatory reporting timeliness",
    ]),
    ("Data & integration governance", [
        "Master data quality (customer, product, inventory, billing account)",
        "Order fallout due to data mismatch",
        "BSS ↔ network ↔ finance reconciliation gaps",
        "Report freshness / data latency",
    ]),
    ("Multivendor ecosystem (managed services)", [
        "Vendor SLA attainment by tower",
        "Cross-vendor escalation time",
        "Shared CPI attainment; penalty exposure",
        "Automation / playbook reuse across vendors",
    ]),
    ("Autonomous operations maturity", [
        "Reactive → predictive → autonomous mix",
        "% incidents auto-remediated (closed-loop)",
        "Top-offender closure rate",
        "SPOG correlation across BSS + OSS + network",
        "Human touch rate on P1 scenarios",
    ]),
]

JOURNEY = [
    ("Buy", "COP/CHQ sales, conversion, fallout", "CRM, digital channels"),
    ("Activate", "O2A, port-in time, provisioning success", "BSS, provisioning, network"),
    ("Use", "Availability, QoS, complaints", "RAN, core, apps"),
    ("Bill", "Cycle time, reruns, accuracy", "Billing, mediation, finance"),
    ("Pay", "Collection stages, DSO, write-off pipeline", "Collections, payments"),
    ("Care", "FCR, repeat contacts, churn save", "CRM, ITSM, network"),
    ("Retain", "Churn from/to by operator, ARPU", "All domains"),
]


def build_docx():
    doc = Document()
    title = doc.add_heading("Telecom Operator KPI Catalogue", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        "Holistic business & operational KPIs for a Tier-1 operator "
        "(BSS + OSS/IT + Network + Infra/Cloud + Security + AO maturity).\n"
        "Prepared by Gaurav Chaudhary · illustrative framework for managed services / control-tower design."
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    doc.add_heading("How to use this", level=1)
    doc.add_paragraph(
        "Build 4 layers: (1) Executive north-stars · (2) Domain scorecards · "
        "(3) Customer journey KPIs · (4) Operational drill-down. "
        "Sample dashboard: operator-kpi-dashboard.html"
    )

    doc.add_heading("Customer journey map", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Journey"
    hdr[1].text = "Business KPIs"
    hdr[2].text = "Underlying stack"
    for j, kpis, stack in JOURNEY:
        row = table.add_row().cells
        row[0].text = j
        row[1].text = kpis
        row[2].text = stack

    for heading, items in SECTIONS:
        doc.add_heading(heading, level=1)
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Your BSS examples — mapped", level=1)
    examples = [
        "Sales across COP & CHQ → Commercial / channel scorecard",
        "Churn (from/to, operator-wise) → Executive + lifecycle",
        "Application availability → OSS/IT + journey (Use)",
        "Order capture & lead time → BSS (Buy/Activate)",
        "Billing E2E, reruns, prep→confirm → BSS (Bill)",
        "Collection stages & write-off queue → Finance (Pay)",
    ]
    for ex in examples:
        doc.add_paragraph(ex, style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph("End-state visual: open operator-kpi-dashboard.html for sample charts.")
    p.runs[0].italic = True

    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")


def wrap(c, text, x, y, max_w, size=10, leading=13):
    c.setFont("Helvetica", size)
    words, lines, cur = text.split(), [], ""
    for w in words:
        test = (cur + " " + w).strip()
        if c.stringWidth(test, "Helvetica", size) <= max_w:
            cur = test
        else:
            if cur:
                lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    for i, line in enumerate(lines):
        c.drawString(x, y - i * leading, line)
    return len(lines) * leading


def build_pdf():
    c = canvas.Canvas(OUT_PDF, pagesize=A4)
    w, h = A4
    margin = 0.65 * inch
    y = h - margin

    c.setFillColor(HexColor("#0F1419"))
    c.rect(0, 0, w, h, fill=1, stroke=0)

    c.setFillColor(HexColor("#D4A853"))
    c.setFont("Helvetica-Bold", 18)
    c.drawString(margin, y, "Telecom Operator KPI Catalogue")
    y -= 22
    c.setFillColor(HexColor("#8B9CB3"))
    c.setFont("Helvetica", 10)
    c.drawString(margin, y, "BSS · OSS/IT · Network · Infra/Cloud · Security · AO maturity")
    y -= 28

    c.setFillColor(HexColor("#E6EDF3"))
    y -= wrap(
        c,
        "Holistic KPI framework for a Tier-1 operator. Four layers: Executive north-stars → Domain scorecards → Customer journey → Operational drill-down. Sample dashboard: operator-kpi-dashboard.html",
        margin, y, w - 2 * margin, size=10,
    ) + 10

    for heading, items in SECTIONS:
        if y < 1.2 * inch:
            c.showPage()
            c.setFillColor(HexColor("#0F1419"))
            c.rect(0, 0, w, h, fill=1, stroke=0)
            y = h - margin

        c.setFillColor(HexColor("#58A6FF"))
        c.setFont("Helvetica-Bold", 11)
        c.drawString(margin, y, heading)
        y -= 16
        c.setFillColor(HexColor("#E6EDF3"))
        c.setFont("Helvetica", 9)
        for item in items[:8]:  # cap per section on PDF for length
            y -= wrap(c, "•  " + item, margin + 8, y, w - 2 * margin - 8, size=9, leading=12) + 2
        if len(items) > 8:
            c.setFillColor(HexColor("#8B9CB3"))
            c.drawString(margin + 8, y, f"  … +{len(items)-8} more (see Word doc)")
            y -= 14
        y -= 8

    c.setFillColor(HexColor("#8B9CB3"))
    c.setFont("Helvetica", 8)
    c.drawString(margin, 0.5 * inch, "Gaurav Chaudhary · operator-kpi-catalog.docx for full list")
    c.save()
    print(f"Wrote {OUT_PDF}")


if __name__ == "__main__":
    build_docx()
    build_pdf()
